#!/usr/bin/env python3
"""
Math Question Generation Script
Generates new math questions similar to base questions using LLM approach
"""

import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import numpy as np
from PIL import Image, ImageDraw, ImageFont
import io

class MathQuestionGenerator:
    def __init__(self):
        self.curriculum = {
            "Quantitative Math": {
                "Problem Solving": ["Numbers and Operations", "Algebra", "Geometry", "Problem Solving", "Probability and Statistics", "Data Analysis"],
                "Algebra": ["Algebraic Word Problems", "Interpreting Variables", "Polynomial Expressions (FOIL/Factoring)", "Rational Expressions", "Exponential Expressions (Product rule, negative exponents)", "Quadratic Equations & Functions (Finding roots/solutions, graphing)", "Functions Operations"],
                "Geometry and Measurement": ["Area & Volume", "Perimeter", "Lines, Angles, & Triangles", "Right Triangles & Trigonometry", "Circles (Area, circumference)", "Coordinate Geometry", "Slope", "Transformations (Dilating a shape)", "Parallel & Perpendicular Lines", "Solid Figures (Volume of Cubes)"],
                "Numbers and Operations": ["Basic Number Theory", "Prime & Composite Numbers", "Rational Numbers", "Order of Operations", "Estimation", "Fractions, Decimals, & Percents", "Sequences & Series", "Computation with Whole Numbers", "Operations with Negatives"],
                "Data Analysis & Probability": ["Interpretation of Tables & Graphs", "Trends & Inferences", "Probability (Basic, Compound Events)", "Mean, Median, Mode, & Range", "Weighted Averages", "Counting & Arrangement Problems"]
            }
        }
    
    def generate_question_1(self):
        """Generate a question similar to the uniform combinations problem"""
        return {
            "title": "Math Assessment - Combinatorics and Counting",
            "description": "This assessment focuses on counting principles and combinatorial problems involving real-world scenarios.",
            "question": "A restaurant offers a lunch special where customers can choose 1 main dish and 1 side dish. The menu shows the following options:\n\n## Lunch Special Menu\n| Main Dish | Side Dish |\n| :---: | :---: |\n| Grilled Chicken | French Fries |\n| Beef Burger | Coleslaw |\n| Fish Fillet | Mashed Potatoes |\n| Vegetarian Pasta | Garden Salad |\n| | Onion Rings |\n\nHow many different lunch combinations are possible?",
            "instruction": "Use the counting principle to determine the total number of possible combinations.",
            "difficulty": "moderate",
            "order": 1,
            "options": [
                "15 combinations",
                "20 combinations", 
                "25 combinations",
                "30 combinations",
                "35 combinations"
            ],
            "correct_answer": "20 combinations",
            "explanation": "Using the counting principle: Number of main dishes × Number of side dishes = 4 × 5 = 20 different combinations. Each main dish can be paired with any of the 5 side dishes, giving us 4 × 5 = 20 total possibilities.",
            "subject": "Quantitative Math",
            "unit": "Problem Solving", 
            "topic": "Counting & Arrangement Problems",
            "plusmarks": 1
        }
    
    def generate_question_2(self):
        """Generate a question similar to the sphere packing problem"""
        return {
            "title": "Math Assessment - Geometry and Spatial Reasoning",
            "description": "This assessment focuses on geometric concepts including area, volume, and spatial relationships.",
            "question": "A cylindrical container is designed to hold 8 tennis balls tightly packed in two layers. Each tennis ball has a radius of 3.5 centimeters. The top view shows a circular arrangement with 4 balls in the bottom layer and 4 balls in the top layer. Which of the following is closest to the dimensions of the cylindrical container?\n\n**Note**: The height accounts for two layers of balls, and the diameter accommodates the circular arrangement.",
            "instruction": "Calculate the dimensions needed to accommodate the tightly packed tennis balls in the cylindrical container.",
            "difficulty": "hard",
            "order": 2,
            "options": [
                "$7 \\times 7 \\times 14$ cm",
                "$7 \\times 7 \\times 21$ cm", 
                "$14 \\times 14 \\times 14$ cm",
                "$14 \\times 14 \\times 21$ cm",
                "$21 \\times 21 \\times 28$ cm"
            ],
            "correct_answer": "$14 \\times 14 \\times 21$ cm",
            "explanation": "Each tennis ball has a radius of 3.5 cm, so diameter = 7 cm. For the circular arrangement of 4 balls, the container diameter must be approximately 14 cm (2 × 7 cm). The height must accommodate 2 layers: 2 × 7 cm = 14 cm, plus some spacing, so approximately 21 cm total. Therefore, the closest dimensions are $14 \\times 14 \\times 21$ cm.",
            "subject": "Quantitative Math",
            "unit": "Geometry and Measurement",
            "topic": "Area & Volume", 
            "plusmarks": 1
        }
    
    def create_question_image_1(self):
        """Create an image for the lunch special menu question"""
        # Create a simple table-like image
        img = Image.new('RGB', (400, 300), color='white')
        draw = ImageDraw.Draw(img)
        
        # Draw table structure
        draw.rectangle([50, 50, 350, 250], outline='black', width=2)
        draw.line([50, 100, 350, 100], fill='black', width=2)  # Header line
        draw.line([200, 50, 200, 250], fill='black', width=2)  # Vertical line
        
        # Draw text (simplified)
        draw.text((75, 65), "Main Dish", fill='black')
        draw.text((225, 65), "Side Dish", fill='black')
        draw.text((75, 115), "Grilled Chicken", fill='black')
        draw.text((225, 115), "French Fries", fill='black')
        draw.text((75, 140), "Beef Burger", fill='black')
        draw.text((225, 140), "Coleslaw", fill='black')
        draw.text((75, 165), "Fish Fillet", fill='black')
        draw.text((225, 165), "Mashed Potatoes", fill='black')
        draw.text((75, 190), "Veg Pasta", fill='black')
        draw.text((225, 190), "Garden Salad", fill='black')
        draw.text((75, 215), "", fill='black')
        draw.text((225, 215), "Onion Rings", fill='black')
        
        return img
    
    def create_question_image_2(self):
        """Create an image for the cylindrical container question"""
        # Create a diagram showing the cylindrical container with tennis balls
        fig, ax = plt.subplots(figsize=(6, 4))
        
        # Draw the cylindrical container
        container_height = 21
        container_radius = 7
        
        # Draw container outline
        circle = plt.Circle((0, 0), container_radius, fill=False, color='black', linewidth=2)
        ax.add_patch(circle)
        
        # Draw tennis balls in circular arrangement
        ball_radius = 3.5
        angles = [0, 90, 180, 270]  # 4 balls in a circle
        for angle in angles:
            x = (container_radius - ball_radius) * np.cos(np.radians(angle))
            y = (container_radius - ball_radius) * np.sin(np.radians(angle))
            ball = plt.Circle((x, y), ball_radius, fill=True, color='yellow', alpha=0.7)
            ax.add_patch(ball)
        
        # Set plot properties
        ax.set_xlim(-10, 10)
        ax.set_ylim(-10, 10)
        ax.set_aspect('equal')
        ax.set_title('Cylindrical Container with Tennis Balls (Top View)')
        ax.grid(True, alpha=0.3)
        
        # Save to bytes
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        buf.seek(0)
        plt.close()
        
        # Convert to PIL Image
        img = Image.open(buf)
        return img
    
    def format_question_output(self, question_data):
        """Format question data according to the specified output format"""
        output = f"""@title {question_data['title']}
@description {question_data['description']}

@question {question_data['question']}
@instruction {question_data['instruction']}
@difficulty {question_data['difficulty']}
@order {question_data['order']}
@option {question_data['options'][0]}
@option {question_data['options'][1]}
@@option {question_data['correct_answer']}
@option {question_data['options'][3]}
@option {question_data['options'][4]}
@explanation {question_data['explanation']}
@subject {question_data['subject']}
@unit {question_data['unit']}
@topic {question_data['topic']}
@plusmarks {question_data['plusmarks']}

"""
        return output
    
    def create_word_document(self, questions):
        """Create a Word document with the generated questions"""
        doc = Document()
        
        # Add title
        title = doc.add_heading('Math Question Generation - Assessment', 0)
        title.alignment = 1  # Center alignment
        
        # Add description
        doc.add_paragraph('This document contains two newly generated math questions following the specified format and curriculum hierarchy.')
        
        # Add questions
        for i, question in enumerate(questions, 1):
            doc.add_heading(f'Question {i}', level=1)
            
            # Add question text
            doc.add_paragraph(question['question'])
            
            # Add options
            doc.add_heading('Options:', level=2)
            for j, option in enumerate(question['options'], 1):
                marker = "✓" if option == question['correct_answer'] else "○"
                doc.add_paragraph(f'{marker} {option}')
            
            # Add explanation
            doc.add_heading('Explanation:', level=2)
            doc.add_paragraph(question['explanation'])
            
            # Add metadata
            doc.add_heading('Question Metadata:', level=2)
            metadata_table = doc.add_table(rows=1, cols=2)
            metadata_table.style = 'Table Grid'
            
            # Add header row
            header_cells = metadata_table.rows[0].cells
            header_cells[0].text = 'Field'
            header_cells[1].text = 'Value'
            
            # Add metadata rows
            metadata_fields = ['Subject', 'Unit', 'Topic', 'Difficulty', 'Marks']
            metadata_values = [
                question['subject'],
                question['unit'], 
                question['topic'],
                question['difficulty'],
                str(question['plusmarks'])
            ]
            
            for field, value in zip(metadata_fields, metadata_values):
                row_cells = metadata_table.add_row().cells
                row_cells[0].text = field
                row_cells[1].text = value
            
            # Add separator
            if i < len(questions):
                doc.add_page_break()
        
        return doc
    
    def generate_all_questions(self):
        """Generate all questions and create outputs"""
        print("Generating new math questions...")
        
        # Generate questions
        question1 = self.generate_question_1()
        question2 = self.generate_question_2()
        questions = [question1, question2]
        
        # Create images
        print("Creating question images...")
        img1 = self.create_question_image_1()
        img2 = self.create_question_image_2()
        
        # Save images
        os.makedirs('images', exist_ok=True)
        img1.save('images/lunch_special_menu.png')
        img2.save('images/tennis_ball_container.png')
        
        # Create formatted output
        print("Creating formatted output...")
        os.makedirs('output', exist_ok=True)
        
        # Save raw formatted output
        with open('output/questions_formatted.txt', 'w') as f:
            for question in questions:
                f.write(self.format_question_output(question))
        
        # Create Word document
        doc = self.create_word_document(questions)
        doc.save('output/Math_Questions_Assessment.docx')
        
        print("Questions generated successfully!")
        print("Output files created in 'output/' directory")
        print("Images saved in 'images/' directory")
        
        return questions

def main():
    """Main function to run the question generator"""
    generator = MathQuestionGenerator()
    questions = generator.generate_all_questions()
    
    print("\n" + "="*50)
    print("GENERATED QUESTIONS SUMMARY")
    print("="*50)
    
    for i, question in enumerate(questions, 1):
        print(f"\nQuestion {i}:")
        print(f"Topic: {question['topic']}")
        print(f"Difficulty: {question['difficulty']}")
        print(f"Correct Answer: {question['correct_answer']}")
        print("-" * 30)

if __name__ == "__main__":
    main() 